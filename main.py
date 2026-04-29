import os
import time
import datetime
from zoneinfo import ZoneInfo
from dotenv import load_dotenv
from groq import Groq
from docx import Document
import smtplib
from email.message import EmailMessage
from googleapiclient.discovery import build
from youtube_transcript_api import YouTubeTranscriptApi

# Load environment variables with explicit encoding to handle potential UTF-8 BOM issue
if os.path.exists(".env"):
    load_dotenv(".env", encoding='utf-8-sig')
else:
    load_dotenv()

# Helper to get and clean environment variables (strip whitespace/BOM just in case)
def get_env_var(name, default=None):
    val = os.getenv(name, default)
    if val and isinstance(val, str):
        return val.strip().replace('\ufeff', '')
    return val

YOUTUBE_API_KEY = get_env_var("YOUTUBE_API_KEY")
GROQ_API_KEY = get_env_var("GROQ_API_KEY")
GMAIL_ADDRESS = get_env_var("GMAIL_ADDRESS")
GMAIL_APP_PASSWORD = get_env_var("GMAIL_APP_PASSWORD")

def get_groq_client():
    """Get the Groq API client."""
    if not GROQ_API_KEY:
        print(f"Error: GROQ_API_KEY environment variable not set. (Current list: {list(os.environ.keys())[:5]}...)")
        return None
    return Groq(api_key=GROQ_API_KEY)
        
def fetch_top_trending_investment_videos():
    """
    Search YouTube for the top trending English investment videos 
    published in the last 48 hours.
    """
    if not YOUTUBE_API_KEY:
        print("Error: YOUTUBE_API_KEY environment variable not set.")
        return None

    youtube = build('youtube', 'v3', developerKey=YOUTUBE_API_KEY)
    
    # Target videos from the last 48 hours. YouTube API requires RFC 3339 formatted date-time value (1970-01-01T00:00:00Z)
    published_after = (datetime.datetime.now(datetime.timezone.utc) - datetime.timedelta(hours=48)).replace(microsecond=0).isoformat().replace('+00:00', 'Z')
    
    try:
        search_response = youtube.search().list(
            q="stock market investing Wall Street finance earnings",
            part="id,snippet",
            maxResults=50,
            order="viewCount",
            publishedAfter=published_after,
            relevanceLanguage="en",
            regionCode="US",
            type="video"
        ).execute()

        if not search_response.get("items"):
            print("No trending investment videos found.")
            return None
            
        return search_response["items"]
        
    except Exception as e:
        print(f"Error fetching YouTube data: {e}")
        return None

def get_video_content(video_id, description_fallback=""):
    """Retrieve the English transcript for a video. Falls back to description if blocked."""
    try:
        transcript_list = YouTubeTranscriptApi().list(video_id)
        transcript = None

        # Step1: 英語（手動・自動生成・地域バリアント）を優先的に探す
        for lang in ['en', 'en-US', 'en-GB', 'en-AU', 'en-CA']:
            try:
                transcript = transcript_list.find_transcript([lang])
                break
            except Exception:
                continue

        # Step2: 英語が見つからない場合、他言語を英語に翻訳して取得
        if transcript is None:
            try:
                available = transcript_list._transcripts
                if available:
                    first_lang = next(iter(available))
                    transcript = transcript_list.find_transcript([first_lang]).translate('en')
                    print(f"   -> No native English transcript. Using translated version from '{first_lang}'.")
            except Exception:
                pass

        if transcript is None:
            raise Exception("No transcript found")

        fetched_transcript = transcript.fetch()
        full_text = " ".join([snippet.text for snippet in fetched_transcript.snippets])
        source = "transcript"

    except Exception as e:
        # トランスクリプト取得失敗 → 説明文にフォールバック
        if description_fallback and description_fallback.strip():
            print(f"   -> Transcript unavailable ({type(e).__name__}). Falling back to video description.")
            full_text = description_fallback.strip()
            source = "description"
        else:
            print(f"   -> Transcript and description both unavailable. Skipping.")
            return None

    max_chars = 10000
    if len(full_text) > max_chars:
        full_text = full_text[:max_chars] + "\n...[Content Truncated]..."

    print(f"   -> Content source: {source} ({len(full_text)} chars)")
    return full_text

def generate_single_video_summary(client, rank, title, channel, transcript):
    """Generate a detailed 1000-character summary for a single video."""
    prompt = f"""
あなたは優秀な証券アナリストです。以下の海外の最新投資YouTube動画のトランスクリプトを読み、日本語で詳細に要約してください。

【重要な指示】
1. この動画１本に対して、要約の文章量が「約1000文字程度」になるように、非常に詳細に解説を展開してください。「具体的な数値」「登場した銘柄名」「市場の背景」「配信者の主張の根拠」などをしっかり盛り込み、読み応えのある構成にしてください。
2. 「風が吹けば桶屋が儲かる」のような独自の将来予測や市場への波及効果の分析は絶対に含めず、あくまで動画内で語られている内容のみを詳細に要約してください。

【対象動画情報】
タイトル: {title}
チャンネル名: {channel}
トランスクリプト: {transcript}

【出力フォーマット】
以下の形式で直接出力してください。その他の挨拶や前置きは一切不要です。（※見出しのフォーマットは厳守）

## 注目トピック 第{rank}位: {title} (チャンネル名: {channel})
（ここに約1000文字の非常に詳細な要約を記述）
"""
    
    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.4,
            max_tokens=2048,
        )
        return chat_completion.choices[0].message.content.strip()
    except Exception as e:
        print(f"Error generating summary for rank {rank}: {e}")
        return f"## 注目トピック 第{rank}位: {title} (チャンネル名: {channel})\n（要約生成中にエラーが発生しました: {e}）"

def generate_filename_summary(client, text_content):
    """Generate a 15-character summary for the filename."""
    prompt = f"""
以下のレポートの内容を一言で表す、15文字以内の簡潔なタイトル（要約）を作成してください。
記号や改行、空白はすべて除き、Windowsのファイル名として使える文字列のみを直接出力してください。
回答には余計な挨拶や「〜です」などは一切含めないでください。
例: 「イラン原油高とAI波及」「エネルギーと技術の波及効果」など

【レポート内容】
{text_content[:1500]}
"""
    try:
        chat_completion = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
            temperature=0.3,
            max_tokens=30,
        )
        summary = chat_completion.choices[0].message.content.strip()
        # Clean up the output string to be safe for filenames
        bad_chars = ['\\', '/', ':', '*', '?', '"', '<', '>', '|', '「', '」', ' ', '　', '\n']
        for c in bad_chars:
            summary = summary.replace(c, '')
        return summary[:15]
    except Exception as e:
        print(f"Error generating summary: {e}")
        return "海外YouTubeトップ10"

def save_as_docx(report_text, filepath):
    """Save the markdown-like text to a nice Word document."""
    doc = Document()
    for line in report_text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            if line.strip():
                doc.add_paragraph(line.strip())
    doc.save(filepath)

def send_email_with_attachment(filepath, filename):
    """Send the generated report via email. Retries up to 3 times on failure."""
    if not GMAIL_ADDRESS or not GMAIL_APP_PASSWORD:
        print("Warning: GMAIL_ADDRESS or GMAIL_APP_PASSWORD not set. Skipping email.")
        return

    msg = EmailMessage()
    msg['Subject'] = f"【自動リサーチ】海外投資YouTubeトップ10要約 ({filename})"
    msg['From'] = GMAIL_ADDRESS
    msg['To'] = GMAIL_ADDRESS
    msg.set_content("本日の海外投資ニュースのリサーチ結果を添付します。\n\n※このメールはGitHub Actionsまたはローカル環境から自動送信されています。")

    try:
        with open(filepath, 'rb') as f:
            file_data = f.read()
            msg.add_attachment(
                file_data, 
                maintype='application', 
                subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', 
                filename=filename
            )
    except Exception as e:
        print(f"Error reading file for attachment: {e}")
        return

    import time
    max_retries = 3
    retry_wait = 30  # seconds

    for attempt in range(1, max_retries + 1):
        try:
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
                smtp.login(GMAIL_ADDRESS, GMAIL_APP_PASSWORD)
                smtp.send_message(msg)
            print("-> Successfully sent email with attached report!")
            return
        except Exception as e:
            print(f"Error sending email (attempt {attempt}/{max_retries}): {e}")
            if attempt < max_retries:
                print(f"   -> Retrying in {retry_wait} seconds...")
                time.sleep(retry_wait)
            else:
                print("   -> All retry attempts failed. Email was not sent.")

def fetch_and_process_news():
    """
    Fetch trending international finance YouTube video and process it for retail investors using Groq.
    """
    jst_time = datetime.datetime.now(ZoneInfo("Asia/Tokyo"))
    date_str = jst_time.strftime("%Y%m%d")
    
    # Save output to specified directory or local 'output' folder
    target_dir = os.getenv("RESEARCH_OUTPUT_DIR", "output")
    os.makedirs(target_dir, exist_ok=True)

    print(f"[{jst_time}] Starting YouTube video research and plot generation...")
    
    client = get_groq_client()
    if not client:
        return
        
    print("1. Fetching top trending English investment YouTube videos...")
    videos = fetch_top_trending_investment_videos()
    
    if not videos:
        print("Stopping process: Could not find any trending videos.")
        return
        
    print("2. Fetching transcripts and filtering for exactly 10 valid videos...")
    collected_videos = []
    valid_count = 0
    
    for idx, item in enumerate(videos, 1):
        video_id = item["id"]["videoId"]
        title = item["snippet"]["title"]
        channel = item["snippet"]["channelTitle"]
        description = item["snippet"].get("description", "")
        
        safe_title = title.encode('cp932', 'replace').decode('cp932')
        print(f"-> Checking original Rank {idx}: {safe_title}")
        
        content_text = get_video_content(video_id, description_fallback=description)
        if not content_text:
            print("   -> No usable content found. Skipping.")
            continue
            
        valid_count += 1
        collected_videos.append({
            "rank": valid_count,
            "title": title,
            "channel": channel,
            "transcript": content_text
        })
        
        if valid_count >= 10:
            print("-> Successfully collected 10 videos with content.")
            break
            
    if valid_count == 0:
        print("Stopping process: Could not find any videos with valid transcripts.")
        return
        
    print("3. Generating detailed 1000-char summary for each video using Groq...")
    report_text = "# 注目トピック　TOP１０\n\n"
    for v in collected_videos:
        print(f"   -> Summarizing Rank {v['rank']}: {v['title'].encode('cp932', 'replace').decode('cp932')}")
        summary_text = generate_single_video_summary(client, v['rank'], v['title'], v['channel'], v['transcript'])
        report_text += summary_text + "\n\n"
        time.sleep(2) # Small delay to avoid API rate limits
    
    print("4. Generating 15-character filename summary...")
    summary = generate_filename_summary(client, report_text)
    
    # Generate filename (e.g., 20260315_01_イラン原油高.docx)
    index = 1
    while True:
        filename = f"{date_str}_{index:02d}_{summary}.docx"
        filepath = os.path.join(target_dir, filename)
        if not os.path.exists(filepath):
            break
        index += 1
        
    print(f"4. Saving report to {filepath} ...")
    try:
        save_as_docx(report_text, filepath)
        print("-> Successfully saved report!")
        
        print("5. Sending report via email...")
        send_email_with_attachment(filepath, filename)
    except Exception as e:
        print(f"Error saving file: {e}")

if __name__ == "__main__":
    fetch_and_process_news()
